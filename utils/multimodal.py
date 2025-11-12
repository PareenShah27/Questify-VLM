"""
utils/multimodal_utils.py - Multimodal Document Processing Utilities

Handles format conversion, document chunking, and hybrid document processing.
Integrates multiple preprocessors for unified multimodal handling.
"""

import logging
from pathlib import Path
from typing import Dict, List, Optional, Tuple, Any, Union
import numpy as np
from dataclasses import dataclass


@dataclass
class Document:
    """Represent a document with metadata."""
    doc_id: str
    content_type: str  # 'text', 'image', 'pdf', 'hybrid'
    file_path: Optional[str] = None
    text_content: Optional[str] = None
    image_data: Optional[np.ndarray] = None
    embedding: Optional[np.ndarray] = None
    metadata: Optional[Dict] = None
    
    def to_dict(self) -> Dict:
        """Convert to dictionary."""
        return {
            'doc_id': self.doc_id,
            'content_type': self.content_type,
            'file_path': self.file_path,
            'text_content': self.text_content,
            'embedding_dim': len(self.embedding) if self.embedding is not None else None,
            'metadata': self.metadata,
        }


class DocumentChunker:
    """Chunk documents for processing."""
    
    def __init__(self, config):
        """Initialize document chunker."""
        self.config = config
        self.logger = logging.getLogger(self.__class__.__name__)
        
        self.chunk_size = config.get('vlm.chunk_size', 512)
        self.stride = config.get('vlm.stride', 256)
    
    def chunk_text(self, text: str, overlap: bool = True) -> List[str]:
        """
        Chunk text into overlapping segments.
        
        Args:
            text: Text to chunk
            overlap: Whether to use overlapping chunks
            
        Returns:
            List of text chunks
        """
        try:
            words = text.split()
            chunks = []
            
            step = self.stride if overlap else self.chunk_size
            
            for i in range(0, len(words), step):
                chunk = ' '.join(words[i:i + self.chunk_size])
                if chunk.strip():
                    chunks.append(chunk)
            
            self.logger.info(f"Chunked text into {len(chunks)} parts")
            return chunks
        except Exception as e:
            self.logger.error(f"Error chunking text: {e}")
            return [text]
    
    def chunk_by_sentences(self, text: str, sentences_per_chunk: int = 5) -> List[str]:
        """Chunk text by sentences."""
        try:
            # Simple sentence splitting
            sentences = text.replace('!', '.').replace('?', '.').split('.')
            sentences = [s.strip() for s in sentences if s.strip()]
            
            chunks = []
            for i in range(0, len(sentences), sentences_per_chunk):
                chunk = ' '.join(sentences[i:i + sentences_per_chunk])
                if chunk.strip():
                    chunks.append(chunk)
            
            return chunks
        except Exception as e:
            self.logger.error(f"Error chunking sentences: {e}")
            return [text]


class FormatConverter:
    """Convert between different document formats."""
    
    def __init__(self, config):
        """Initialize format converter."""
        self.config = config
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def markdown_to_text(self, markdown: str) -> str:
        """Convert markdown to plain text."""
        try:
            import re
            # Remove markdown formatting
            text = re.sub(r'[*_~`]', '', markdown)
            text = re.sub(r'#+\s+', '', text)
            text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
            text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', '', text)
            return text.strip()
        except Exception as e:
            self.logger.error(f"Error converting markdown: {e}")
            return markdown
    
    def docx_to_text(self, docx_path: str) -> Optional[str]:
        """Extract text from DOCX file."""
        try:
            # Import the top-level module to avoid relying on resolving the 'Document' symbol
            from docx import Document
            
            doc = Document(docx_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return text
        except ImportError:
            self.logger.warning("python-docx not installed. Cannot read DOCX.")
            return None
        except Exception as e:
            self.logger.error(f"Error reading DOCX: {e}")
            return None
    
    def json_to_text(self, json_data: Dict) -> str:
        """Convert JSON to readable text."""
        try:
            import json
            lines = []
            
            def extract_text(obj, indent=0):
                if isinstance(obj, dict):
                    for k, v in obj.items():
                        lines.append('  ' * indent + f"{k}:")
                        extract_text(v, indent + 1)
                elif isinstance(obj, list):
                    for item in obj:
                        extract_text(item, indent + 1)
                else:
                    lines.append('  ' * indent + str(obj))
            
            extract_text(json_data)
            return '\n'.join(lines)
        except Exception as e:
            self.logger.error(f"Error converting JSON: {e}")
            return str(json_data)


class FileHandler:
    """Handle various file types."""
    
    def __init__(self, config):
        """Initialize file handler."""
        self.config = config
        self.logger = logging.getLogger(self.__class__.__name__)
    
    def read_text_file(self, file_path: str, encoding: str = 'utf-8') -> Optional[str]:
        """Read text file."""
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except Exception as e:
            self.logger.error(f"Error reading file: {e}")
            return None
    
    def write_text_file(self, file_path: str, content: str, encoding: str = 'utf-8') -> bool:
        """Write text file."""
        try:
            Path(file_path).parent.mkdir(parents=True, exist_ok=True)
            with open(file_path, 'w', encoding=encoding) as f:
                f.write(content)
            return True
        except Exception as e:
            self.logger.error(f"Error writing file: {e}")
            return False
    
    def get_file_info(self, file_path: Union[str, Path]) -> Optional[Dict[str, Any]]:
        """Get file information."""
        try:
            file_path = Path(file_path)
            
            if not file_path.exists():
                return None
            
            return {
                'name': file_path.name,
                'size_bytes': file_path.stat().st_size,
                'extension': file_path.suffix,
                'created_at': file_path.stat().st_ctime,
                'modified_at': file_path.stat().st_mtime,
            }
        except Exception as e:
            self.logger.error(f"Error getting file info: {e}")
            return None
    
    def batch_read_files(self, file_paths: List[str]) -> Dict[str, Optional[str]]:
        """Read multiple files."""
        results = {}
        for file_path in file_paths:
            results[file_path] = self.read_text_file(file_path)
        return results


class DocumentProcessor:
    """High-level document processing coordinator."""
    
    def __init__(self, config, text_preprocessor, image_preprocessor):
        """Initialize document processor."""
        self.config = config
        self.logger = logging.getLogger(self.__class__.__name__)
        
        self.text_preprocessor = text_preprocessor
        self.image_preprocessor = image_preprocessor
        self.chunker = DocumentChunker(config)
        self.converter = FormatConverter(config)
        self.file_handler = FileHandler(config)
    
    def process_text_document(self, text: str, doc_id: str) -> Document:
        """Process text document."""
        try:
            processed_text = self.text_preprocessor.preprocess(text)
            
            return Document(
                doc_id=doc_id,
                content_type='text',
                text_content=processed_text,
                metadata={'original_length': len(text)},
            )
        except Exception as e:
            self.logger.error(f"Error processing text: {e}")
            return Document(doc_id=doc_id, content_type='text')
    
    def process_image_document(self, image_path: str, doc_id: str) -> Document:
        """Process image document."""
        try:
            image = self.image_preprocessor.load_image(image_path)
            
            if image is None:
                return Document(doc_id=doc_id, content_type='image')
            
            image_array = self.image_preprocessor.normalize_image(image)
            stats = self.image_preprocessor.get_image_statistics(image)
            
            return Document(
                doc_id=doc_id,
                content_type='image',
                file_path=str(image_path),
                image_data=image_array,
                metadata=stats,
            )
        except Exception as e:
            self.logger.error(f"Error processing image: {e}")
            return Document(doc_id=doc_id, content_type='image')
    
    def process_hybrid_document(self, file_path: Union[str, Path], doc_id: str) -> Document:
        """Process document with both text and visual components."""
        try:
            file_path = Path(file_path)
            
            # For PDFs, extract both text and images
            if file_path.suffix.lower() == '.pdf':
                from .image_preprocessor import PDFProcessor
                pdf_processor = PDFProcessor(self.config)
                
                images = pdf_processor.pdf_to_images(str(file_path))
                pdf_info = pdf_processor.get_pdf_info(str(file_path))
                
                return Document(
                    doc_id=doc_id,
                    content_type='pdf',
                    file_path=str(file_path),
                    metadata=pdf_info,
                )
            
            return Document(doc_id=doc_id, content_type='unknown')
        except Exception as e:
            self.logger.error(f"Error processing hybrid document: {e}")
            return Document(doc_id=doc_id, content_type='unknown')